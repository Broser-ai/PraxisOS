"""
Genererer PraxisOS-Forarbejde.docx — komplet forarbejde-liste for Michael.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path


INK = RGBColor(0x1B, 0x1A, 0x17)
MUTED = RGBColor(0x55, 0x52, 0x4D)
ACCENT = RGBColor(0x8A, 0x6A, 0x3D)


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = INK
        run.font.name = "Calibri"
    return h


def add_para(doc, text, bold=False, italic=False, size=11, color=INK):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.name = "Calibri"
    r.font.color.rgb = color
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.runs[0] if p.runs else p.add_run(text)
    if not p.runs:
        r = p.add_run(text)
    else:
        p.runs[0].text = text
    p.runs[0].font.size = Pt(11)
    p.runs[0].font.name = "Calibri"
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"

    # Header
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr[i], "1B1A17")

    # Rows
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10.5)

    if col_widths:
        for ri in range(len(table.rows)):
            for ci, w in enumerate(col_widths):
                table.rows[ri].cells[ci].width = Cm(w)


# ============================================================
doc = Document()

# Margins
for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

# Default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ------------------------------------------------------------
# TITEL
# ------------------------------------------------------------
title = doc.add_heading("PraxisOS · Forarbejde til go-live", level=0)
for r in title.runs:
    r.font.color.rgb = INK

sub = doc.add_paragraph()
r = sub.add_run("Komplet checkliste for kommerciel lancering · juni 2026")
r.italic = True
r.font.size = Pt(12)
r.font.color.rgb = MUTED

doc.add_paragraph()  # spacer

# ------------------------------------------------------------
# STATUS
# ------------------------------------------------------------
add_heading(doc, "Status pr. 2026-06-15", level=1)
add_para(doc,
    "Koden er ~85% færdig. Build er clean, 43/43 ruter returnerer 200, "
    "by Pilar er markeret som trial-tenant (gratis ubegrænset), public landing + "
    "signup-flow + pricing er live på localhost. De sidste 15% kræver eksterne "
    "aftaler (Supabase EU, Stripe, MitID broker m.fl.) før de kan kodes færdigt.")

doc.add_paragraph()

# ============================================================
# 1 · KONTI
# ============================================================
add_heading(doc, "1 · Konti du selv skal oprette (kan klares i dag)", level=1)
add_table(doc,
    ["#", "Konto", "Hvor", "Hvad du skal gøre", "Tid", "Pris"],
    [
        ["1.1", "Supabase", "supabase.com",
         "Opret org 'PraxisOS' · projekt i eu-central-1 (Frankfurt) · gem URL + anon key + service-role key",
         "10 min", "gratis tier OK"],
        ["1.2", "Vercel", "vercel.com",
         "Opret team 'PraxisOS' · connect GitHub-repo",
         "5 min", "gratis tier OK"],
        ["1.3", "GitHub", "github.com",
         "Opret privat repo 'praxisos' · push prototypen",
         "5 min", "gratis"],
        ["1.4", "Sentry", "sentry.io",
         "Opret projekt Next.js · gem DSN",
         "5 min", "gratis · 5k events/md"],
        ["1.5", "OpenAI Platform", "platform.openai.com",
         "Opret API-key med billing-cap (start 50 USD)",
         "5 min", "pay-as-you-go"],
        ["1.6", "Domæne", "unoeuro.dk eller GoDaddy",
         "Reservér praxisos.dk + praxis.app (+ evt. praxisos.io)",
         "10 min", "~80 kr/år/stk"],
    ],
    col_widths=[0.8, 2.5, 3.0, 6.0, 1.5, 2.5])

doc.add_paragraph()
add_para(doc, "Send mig efter dette:", bold=True)
add_bullet(doc, "Supabase: URL + anon key + service-role key")
add_bullet(doc, "Vercel: project-id eller team-invite som developer")
add_bullet(doc, "GitHub: repo-URL")
add_bullet(doc, "Domæner: hvilke du fik reserveret")

doc.add_paragraph()

# ============================================================
# 2 · CVR + selskab
# ============================================================
add_heading(doc, "2 · CVR + selskab (skal være på plads før kunder)", level=1)
add_table(doc,
    ["#", "Hvad", "Hvor", "Tid"],
    [
        ["2.1", "PraxisOS ApS stiftet", "virk.dk", "1-3 dage"],
        ["2.2", "Erhvervskonto i bank + IBAN", "din bank", "1-2 uger"],
        ["2.3", "Momsregistreret + SE-nummer", "skat.dk", "inkl. i stiftelse"],
        ["2.4", "D&B nummer (DUNS)", "dnb.com", "14 dage · gratis"],
    ],
    col_widths=[0.8, 5.5, 5.5, 4.5])

doc.add_paragraph()
add_para(doc,
    "Hvis selskabet allerede findes: send mig CVR-nummeret. Jeg opdaterer lib/payments.ts "
    "og signup-flow så det rigtige selskab vises i stedet for placeholderen 99887766.",
    italic=True, color=MUTED)

doc.add_paragraph()

# ============================================================
# 3 · E-SIGNERING med MitID — NYT
# ============================================================
add_heading(doc, "3 · E-signering af kontrakter med MitID", level=1)
add_para(doc,
    "Vi skal kunne lade klinikker underskrive PraxisOS-kontrakten elektronisk med MitID "
    "(både MitID Privat for solister og MitID Erhverv for selskaber). "
    "Her er prissammenligning af de relevante udbydere:")

doc.add_paragraph()

add_table(doc,
    ["Udbyder", "Pris-model", "Per MitID-signatur", "Månedlig fast", "Bemærkning"],
    [
        ["Criipto Signatures", "pay-per-use", "~4-6 kr", "0 kr",
         "BILLIGST · samme broker vi bruger til MitID-login · ingen abonnement"],
        ["eSignatur.dk", "abonnement", "inkl. i pakke", "fra 99 kr/md",
         "Billigst hvis < 30 sign./md"],
        ["Dokobit", "hybrid", "~5 kr", "~99 kr/md", "MitID via Signicat"],
        ["Visma Addo", "abonnement", "inkl.", "fra 199 kr/md", "Stort i DK SMB"],
        ["Penneo", "abonnement", "inkl.", "fra 199 kr/md", "Stort i advokat/revisor"],
        ["Scrive", "abonnement", "inkl.", "fra 249 kr/md", "Svensk · MitID OK"],
    ],
    col_widths=[3.0, 2.5, 3.0, 2.5, 5.5])

doc.add_paragraph()
add_para(doc, "Anbefaling: Criipto Signatures", bold=True)
add_bullet(doc, "Vi bruger allerede Criipto / Signaturgruppen som MitID-broker — samme aftale dækker")
add_bullet(doc, "Pay-per-use → 0 kr månedligt loft mens vi tester med få kunder")
add_bullet(doc, "Ved 50 nye klinikker/år → ca. 200-300 kr/år total (vs. 2.400+ kr hos Scrive)")
add_bullet(doc, "Ved 50+ signaturer/md vurderes skift til Visma Addo eller eSignatur for forudsigelig pris")

doc.add_paragraph()
add_para(doc, "Action: Skriv til Criipto (sales@criipto.com) — bed om at få Signatures aktiveret oven på den eksisterende MitID-broker-aftale (samme som under afsnit 4.1 nedenfor).",
         italic=True, color=ACCENT)

doc.add_paragraph()

# ============================================================
# 4 · DK-integrationer
# ============================================================
add_heading(doc, "4 · DK-integrationer · ansøg parallelt (lead-time 2-8 uger)", level=1)
add_para(doc, "Disse skal startes NU så de er klar når koden er klar.")

doc.add_paragraph()

# 4.1
add_heading(doc, "4.1 · MitID broker · Signaturgruppen (~2 uger)", level=2)
add_bullet(doc, "Kontakt: support@signaturgruppen.dk")
add_bullet(doc, "Du skal: CVR + selskabsnavn + brug-case (klinisk SaaS · NSIS Substantial) · bed om OIDC client_id og test-environment")
add_bullet(doc, "Du modtager: client_id, client_secret, redirect_uri-whitelist, sandbox-credentials")
add_bullet(doc, "Pris: ~3-5 kr per login produktion · gratis test")
add_bullet(doc, "Send mig: alle credentials når du har dem")

# 4.2
add_heading(doc, "4.2 · Trustaftale · Sundhedsdatastyrelsen (~6 uger)", level=2)
add_bullet(doc, "Hvor: nspop.dk · ansøg ledelses-erklæring + ISO 27001-light")
add_bullet(doc, "Du skal: Ansøge om trustaftale for FMK + Min Log · 5-trins onboarding")
add_bullet(doc, "Modtager: trust-certifikat + NSP-client-cert")
add_bullet(doc, "Pris: gratis (men kræver pen-test, ca. 25-40k kr engangs)")
add_bullet(doc, "Send mig: PEM-certifikater når de kommer")

# 4.3
add_heading(doc, "4.3 · MedCom EAN + VANS (~8 uger)", level=2)
add_bullet(doc, "Kontakt: medcom@medcom.dk")
add_bullet(doc, "Du skal: Ansøg om 13-cifret EAN-adresse · vælg VANS-provider (KMD, EG, NNIT)")
add_bullet(doc, "Pris: EAN ~0 kr · VANS ~1-3 kr per besked")
add_bullet(doc, "Send mig: EAN-adresse + VANS-credentials")

# 4.4
add_heading(doc, "4.4 · NemSMS · KOMBIT (~3-4 uger)", level=2)
add_bullet(doc, "Kontakt: kombit.dk")
add_bullet(doc, "Du skal: Ansøg om sender-id (f.eks. 'PraxisOS')")
add_bullet(doc, "Pris: ~0,18 kr/sms")
add_bullet(doc, "Send mig: sender-id + API-credentials")

# 4.5
add_heading(doc, "4.5 · Sygeforsikringen 'danmark' (~4-6 uger)", level=2)
add_bullet(doc, "Kontakt: sygeforsikring.dk · partner-aftale")
add_bullet(doc, "Du skal: Ansøg om webservice-aftale (EDIFACT D04A)")
add_bullet(doc, "Pris: gratis")
add_bullet(doc, "Send mig: partner-id + API-key")

# 4.6
add_heading(doc, "4.6 · Acquiring-bank for PraxisOS Pay", level=2)
add_para(doc, "Valgmulighed A · gå live hurtigt (anbefalet):", bold=True)
add_bullet(doc, "Stripe Connect Custom under-the-hood — 2-3 dages onboarding, alle betalingsmetoder klar")
add_bullet(doc, "Stripe-konto: dashboard.stripe.com · pris 1,4% + 1,80 kr per transaktion")

add_para(doc, "Valgmulighed B · egen acquiring:", bold=True)
add_bullet(doc, "Direkte aftale med Nets/Adyen/Worldline (~3-6 måneder onboarding)")
add_bullet(doc, "Billigere på sigt, langsomt at komme i gang")

add_para(doc, "Min anbefaling: Start med A. Skift til B når omsætning > 5M kr/år.",
         italic=True, color=ACCENT)

doc.add_paragraph()

# ============================================================
# 5 · Compliance + jura
# ============================================================
add_heading(doc, "5 · Compliance · jura (kan parallel-køres)", level=1)
add_table(doc,
    ["#", "Dokument", "Hvem laver det", "Tid"],
    [
        ["5.1", "DPA-skabelon (databehandleraftale)", "Du · tilpas standard fra Datatilsynet", "1 uge"],
        ["5.2", "Sub-processor liste (Supabase, Vercel, OpenAI, Sentry, Stripe, Criipto)",
         "Jeg lister, du verificerer DPAs er signérbar", "1 dag"],
        ["5.3", "Art. 30-fortegnelse", "Jeg laver skabelon, du udfylder", "1 dag"],
        ["5.4", "Privatlivspolitik + Cookie-politik", "Jeg skriver udkast i da-DK", "1 dag"],
        ["5.5", "Servicevilkår for kunder", "Du m. advokat (ca. 8-15k kr)", "2 uger"],
    ],
    col_widths=[0.8, 6.0, 5.5, 2.0])

doc.add_paragraph()
add_para(doc,
    "Min anbefaling: Brug en advokat med GDPR + health-tech erfaring. "
    "Plesner, Bird&Bird DK, eller en mindre boutique. "
    "~15-25k kr engangs for det hele.",
    italic=True, color=MUTED)

doc.add_paragraph()

# ============================================================
# 6 · Pen-test
# ============================================================
add_heading(doc, "6 · Pen-test + ISO 27001-light (~4 uger)", level=1)
add_bullet(doc, "Kontakt: nccgroup.com (Aalborg-kontor) eller Improsec")
add_bullet(doc, "Pris: ~30-50k kr engangs")
add_bullet(doc, "Output: rapport + tilrettelser")
add_bullet(doc, "Lead-time: Book NU · de har 4-8 ugers ventetid")

doc.add_paragraph()

# ============================================================
# 7 · Hvad jeg laver parallelt
# ============================================================
add_heading(doc, "7 · Hvad jeg laver i koden parallelt", level=1)

add_heading(doc, "Sprint 1 · Når Supabase + Vercel er oprettet (1-2 dage)", level=2)
add_bullet(doc, "/api/signup POST-handler — opretter rigtig tenant i Supabase")
add_bullet(doc, "lib/supabase.ts server-client — bytter mock-data ud med rigtig DB")
add_bullet(doc, "Migrations push — kører 0001_initial_schema.sql mod prod")
add_bullet(doc, "Vercel deploy konfiguration — vercel.ts + env-vars")
add_bullet(doc, "First deploy til midlertidigt praxisos.vercel.app")

add_heading(doc, "Sprint 2 · Faktura + billing-engine (3-4 dage)", level=2)
add_bullet(doc, "Stripe Connect setup (når du har Stripe-konto)")
add_bullet(doc, "Subscription-management — månedlig opkrævning per plan")
add_bullet(doc, "Usage-based billing — per-brug fees (MitID, NemSMS, Niels per session)")
add_bullet(doc, "Faktura-PDF generator (dansk format, momsfri sundhedsydelser)")
add_bullet(doc, "/admin/plan/faktura-historik — kundens faktura-arkiv")
add_bullet(doc, "Email-pipeline (booking-bekræftelse, faktura, påmindelser)")

add_heading(doc, "Sprint 3 · Kontrakt-signering med MitID via Criipto (1 dag)", level=2)
add_bullet(doc, "Generér PraxisOS-tenant-kontrakt PDF dynamisk fra signup-data")
add_bullet(doc, "POST kontrakt-hash til Criipto Signatures · MitID Erhverv")
add_bullet(doc, "Modtag signed-PDF webhook · gem i Supabase Storage")
add_bullet(doc, "Send kopier via email (PraxisOS + kunde)")
add_bullet(doc, "Vis underskrevet kontrakt i /admin/contract")

add_heading(doc, "Sprint 4 · Atlas (sidste agent) (1-2 dage)", level=2)
add_bullet(doc, "Atlas self-reflecting engine — kode-gen + selv-reflection-loop")
add_bullet(doc, "Eksempel-flow: kundens ønske → Atlas genererer modul → review → deploy")

add_heading(doc, "Sprint 5 · Polish til salg (1-2 dage)", level=2)
add_bullet(doc, "Onboarding-tour for nye tenants (4-trins guide første gang)")
add_bullet(doc, "Marketing-side udvidet med kunde-testimonials (kan starte med Pilar)")
add_bullet(doc, "Demo-video embed på landing")
add_bullet(doc, "Help-center (FAQ + video-tutorials placeholders)")

doc.add_paragraph()

# ============================================================
# 8 · Tidsplan
# ============================================================
add_heading(doc, "8 · Anbefalet rækkefølge for dig", level=1)

add_heading(doc, "I dag", level=2)
add_bullet(doc, "Opret Supabase + Vercel + GitHub + Sentry + OpenAI konti (~45 min)")
add_bullet(doc, "Reservér domæner")
add_bullet(doc, "Send mig credentials → jeg deployer første version inden for 24t")

add_heading(doc, "I denne uge", level=2)
add_bullet(doc, "Stift / bekræft selskabet")
add_bullet(doc, "Skriv til Criipto om Signatures-tier (e-signering)")
add_bullet(doc, "Skriv til Signaturgruppen (MitID broker), MedCom, KOMBIT, sygeforsikring.dk — alle samtidigt")
add_bullet(doc, "Book pen-test og advokat")

add_heading(doc, "Næste 2 uger", level=2)
add_bullet(doc, "Stripe Connect-konto + send mig keys → jeg builder billing")
add_bullet(doc, "DPA-skabelon på plads")
add_bullet(doc, "Modtag MitID test-credentials → jeg kobler dem på")
add_bullet(doc, "Aktivér Criipto Signatures → jeg kobler kontrakt-flow")

add_heading(doc, "Næste 4-6 uger", level=2)
add_bullet(doc, "Onboard MedCom, FMK, NemSMS efterhånden som aftaler kommer hjem")

doc.add_paragraph()

# ============================================================
# 9 · CHECKLIST
# ============================================================
add_heading(doc, "9 · Send mig disse informationer når du har dem", level=1)

checklist = [
    "CVR-nummer for PraxisOS ApS",
    "Supabase URL + anon-key + service-role-key",
    "Vercel team-invite eller token",
    "GitHub repo-URL",
    "Sentry DSN",
    "OpenAI API-key",
    "Domæne(r) registreret",
    "Stripe API-keys (når oprettet)",
    "Criipto Signatures API-key (når aktiveret)",
    "MitID client_id + secret (når Signaturgruppen leverer)",
    "MedCom EAN (når MedCom svarer)",
    "NemSMS sender-id (når KOMBIT godkender)",
    "Sygesikringen partner-id (når aftale signeret)",
    "Sundhedsdatastyrelsen trust-cert (efter 6 uger)",
]
for item in checklist:
    p = doc.add_paragraph()
    r = p.add_run(f"[ ]  {item}")
    r.font.name = "Consolas"
    r.font.size = Pt(10.5)

doc.add_paragraph()

# ============================================================
# 10 · Kritisk-sti
# ============================================================
add_heading(doc, "10 · Kritisk-sti", level=1)
p = doc.add_paragraph()
r = p.add_run("Supabase + Vercel + Stripe + selskab + Criipto Signatures = ")
r.font.size = Pt(11)
r2 = p.add_run("du kan begynde at sælge inden for 7 dage.")
r2.bold = True
r2.font.size = Pt(11)
r2.font.color.rgb = ACCENT

add_para(doc,
    "Resten (MitID broker, MedCom, FMK, sygesikringen) tager 4-8 uger men kan kobles "
    "på løbende uden at stoppe salget. by Pilar (pilot-kunde) kører gratis i mellemtiden.",
    italic=True, color=MUTED)

doc.add_paragraph()
doc.add_paragraph()

# Sidste linje
p = doc.add_paragraph()
r = p.add_run("Generet 2026-06-15 · PraxisOS prototype · ma@keap.me")
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = MUTED
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ------------------------------------------------------------
# Gem
# ------------------------------------------------------------
out = Path(__file__).resolve().parent.parent / "PraxisOS-Forarbejde.docx"
doc.save(out)
print(f"Skrevet: {out}")
print(f"Størrelse: {out.stat().st_size:,} bytes")
